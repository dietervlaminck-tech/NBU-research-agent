"""Document exports: transcripts (docx/md), articles (docx/md/html/latex),
literature reviews (md/docx/bibtex)."""
import io
import re
from html import escape

import markdown as md_lib
from docx import Document

from ... import db
from .common import slugify, speaker, transcript_text

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# --- minimal markdown -> docx -------------------------------------------------

_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _add_runs(par, text):
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = par.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        else:
            par.add_run(part)


def md_to_docx(doc, md_text):
    """Render markdown headings/lists/bold/italic into a python-docx Document."""
    for raw in (md_text or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2).strip("# ").strip(), level=min(len(m.group(1)), 4))
            continue
        m = re.match(r"^\s*[-*+]\s+(.*)", line)
        if m:
            _add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            continue
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            _add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            continue
        _add_runs(doc.add_paragraph(), line)


def _docx_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- minimal markdown -> latex ------------------------------------------------

def _latex_escape(s):
    s = s.replace("\\", "\x00")
    s = s.replace("{", r"\{").replace("}", r"\}")
    s = s.replace("\x00", r"\textbackslash{}")
    for ch, rep in (("&", r"\&"), ("%", r"\%"), ("$", r"\$"),
                    ("#", r"\#"), ("_", r"\_")):
        s = s.replace(ch, rep)
    s = s.replace("~", r"\textasciitilde{}").replace("^", r"\textasciicircum{}")
    return s


def _latex_inline(s):
    s = _latex_escape(s)
    s = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", s)
    s = re.sub(r"\*(.+?)\*", r"\\textit{\1}", s)
    s = re.sub(r"`(.+?)`", r"\\texttt{\1}", s)
    return s


_LATEX_HEADINGS = {1: "section", 2: "subsection", 3: "subsubsection", 4: "paragraph"}


def md_to_latex_body(md_text):
    out, list_env = [], None

    def close_list():
        nonlocal list_env
        if list_env:
            out.append(f"\\end{{{list_env}}}")
            list_env = None

    for raw in (md_text or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            close_list()
            out.append("")
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            close_list()
            cmd = _LATEX_HEADINGS.get(min(len(m.group(1)), 4), "paragraph")
            out.append(f"\\{cmd}{{{_latex_inline(m.group(2))}}}")
            continue
        m = re.match(r"^\s*[-*+]\s+(.*)", line)
        if m:
            if list_env != "itemize":
                close_list()
                out.append("\\begin{itemize}")
                list_env = "itemize"
            out.append(f"  \\item {_latex_inline(m.group(1))}")
            continue
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            if list_env != "enumerate":
                close_list()
                out.append("\\begin{enumerate}")
                list_env = "enumerate"
            out.append(f"  \\item {_latex_inline(m.group(1))}")
            continue
        close_list()
        out.append(_latex_inline(line))
    close_list()
    return "\n".join(out)


# --- study transcript exports (interviews) ------------------------------------

def _study_sessions(study_id):
    study = db.get("studies", study_id) or {}
    sessions = db.query("sessions", "study_id = ?",
                        (study.get("id", study_id),), order="started_at")
    return study, sessions


def study_docx(study_id):
    study, sessions = _study_sessions(study_id)
    doc = Document()
    doc.add_heading(study.get("title") or "Interview study", level=0)
    if study.get("research_question"):
        doc.add_paragraph(study["research_question"]).runs[0].italic = True
    for i, s in enumerate(sessions, 1):
        doc.add_heading(s.get("respondent_name") or f"Session {i}", level=1)
        meta = doc.add_paragraph()
        meta.add_run(
            f"Session {s.get('id', '')} — started {s.get('started_at', '')}"
            + (f", completed {s['completed_at']}" if s.get("completed_at") else "")
        ).italic = True
        for msg in s.get("messages") or []:
            if not isinstance(msg, dict):
                continue
            p = doc.add_paragraph()
            p.add_run(f"{speaker(msg.get('role'))}: ").bold = True
            p.add_run(msg.get("content", ""))
    return _docx_bytes(doc), f"{slugify(study.get('title'))}-transcripts.docx", DOCX_MIME


def study_md(study_id):
    study, sessions = _study_sessions(study_id)
    parts = [f"# {study.get('title') or 'Interview study'}", ""]
    if study.get("research_question"):
        parts += [f"*{study['research_question']}*", ""]
    for i, s in enumerate(sessions, 1):
        name = s.get("respondent_name") or f"Session {i}"
        parts += [f"## {name} ({s.get('started_at', '')})", ""]
        for msg in s.get("messages") or []:
            if not isinstance(msg, dict):
                continue
            parts += [f"**{speaker(msg.get('role'))}:** {msg.get('content', '')}", ""]
    data = "\n".join(parts).encode("utf-8")
    return data, f"{slugify(study.get('title'))}-transcripts.md", "text/markdown"


# --- article exports ------------------------------------------------------------

def _article_content(article):
    return article.get("content_md") or f"# {article.get('title', '')}\n"


def article_md(article_id):
    article = db.get("articles", article_id) or {}
    data = _article_content(article).encode("utf-8")
    return data, f"{slugify(article.get('title'))}.md", "text/markdown"


def article_html(article_id):
    article = db.get("articles", article_id) or {}
    body = md_lib.markdown(_article_content(article), extensions=["extra"])
    html = (
        "<!DOCTYPE html>\n<html lang=\"en\">\n<head>\n<meta charset=\"utf-8\">\n"
        f"<title>{escape(article.get('title', ''))}</title>\n"
        "<style>body{font-family:Georgia,serif;max-width:46em;margin:2em auto;"
        "padding:0 1em;line-height:1.6}</style>\n</head>\n<body>\n"
        f"{body}\n</body>\n</html>\n"
    )
    return html.encode("utf-8"), f"{slugify(article.get('title'))}.html", "text/html"


def article_docx(article_id):
    article = db.get("articles", article_id) or {}
    doc = Document()
    doc.add_heading(article.get("title") or "Article", level=0)
    metadata = article.get("metadata") or {}
    if metadata.get("author"):
        doc.add_paragraph(str(metadata["author"])).runs[0].italic = True
    md_to_docx(doc, article.get("content_md", ""))
    return _docx_bytes(doc), f"{slugify(article.get('title'))}.docx", DOCX_MIME


def article_latex(article_id):
    article = db.get("articles", article_id) or {}
    metadata = article.get("metadata") or {}
    author = str(metadata.get("author") or metadata.get("authors") or "")
    tex = "\n".join([
        "\\documentclass{article}",
        "\\usepackage[utf8]{inputenc}",
        f"\\title{{{_latex_escape(article.get('title', ''))}}}",
        f"\\author{{{_latex_escape(author)}}}",
        "\\date{\\today}",
        "",
        "\\begin{document}",
        "\\maketitle",
        "",
        md_to_latex_body(article.get("content_md", "")),
        "",
        "\\end{document}",
        "",
    ])
    return tex.encode("utf-8"), f"{slugify(article.get('title'))}.tex", "application/x-tex"


# --- literature review exports --------------------------------------------------

def _review_with_sources(review_id):
    review = db.get("literature_reviews", review_id) or {}
    sources = db.query("sources", "review_id = ?", (review.get("id", review_id),))
    return review, sources


def _source_line(s):
    bits = []
    if s.get("authors"):
        bits.append(s["authors"])
    if s.get("year"):
        bits.append(f"({s['year']})")
    bits.append(s.get("title", ""))
    if s.get("venue"):
        bits.append(s["venue"])
    if s.get("doi"):
        bits.append(f"https://doi.org/{s['doi']}")
    elif s.get("url"):
        bits.append(s["url"])
    return ". ".join(b.strip(". ") for b in bits if b) + "."


def review_md(review_id):
    review, sources = _review_with_sources(review_id)
    parts = [review.get("report_md") or f"# {review.get('research_question', '')}\n"]
    if sources:
        parts += ["", "## References", ""]
        parts += [f"- {_source_line(s)}" for s in sources]
    name = slugify(review.get("research_question"), "literature-review")
    return "\n".join(parts).encode("utf-8"), f"{name}.md", "text/markdown"


def review_docx(review_id):
    review, sources = _review_with_sources(review_id)
    doc = Document()
    report = review.get("report_md", "")
    if not report.lstrip().startswith("#"):
        doc.add_heading(review.get("research_question") or "Literature review", level=0)
    md_to_docx(doc, report)
    if sources:
        doc.add_heading("References", level=1)
        for s in sources:
            doc.add_paragraph(_source_line(s))
    name = slugify(review.get("research_question"), "literature-review")
    return _docx_bytes(doc), f"{name}.docx", DOCX_MIME


def _bibtex_key(source, used):
    first = re.split(r"\s*(?:,|;|&| and )\s*", source.get("authors") or "")[0].strip()
    surname = first.split()[-1] if first else ""
    base = re.sub(r"[^A-Za-z0-9]", "", surname).lower() + (source.get("year") or "")
    base = base or "source"
    if base not in used:
        used.add(base)
        return base
    for i in range(26 * 27):  # a..z, aa..zz
        suffix = (chr(97 + i // 26 - 1) if i >= 26 else "") + chr(97 + i % 26)
        key = base + suffix
        if key not in used:
            used.add(key)
            return key
    key = base + str(len(used))
    used.add(key)
    return key


def _bib_escape(s):
    return str(s).replace("{", "").replace("}", "").strip()


def review_bibtex(review_id):
    review, sources = _review_with_sources(review_id)
    used, entries = set(), []
    for s in sources:
        fields = [("title", s.get("title")), ("author", s.get("authors")),
                  ("year", s.get("year")), ("journal", s.get("venue")),
                  ("doi", s.get("doi")), ("url", s.get("url"))]
        lines = [f"@article{{{_bibtex_key(s, used)},"]
        lines += [f"  {k} = {{{_bib_escape(v)}}}," for k, v in fields if v]
        lines.append("}")
        entries.append("\n".join(lines))
    data = ("\n\n".join(entries) + "\n").encode("utf-8")
    name = slugify(review.get("research_question"), "literature-review")
    return data, f"{name}.bib", "application/x-bibtex"
